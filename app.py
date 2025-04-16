import streamlit as st
import pandas as pd
import random
import io
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Streamlit Page Config
st.set_page_config(page_title="Random Question Generator", page_icon="❓", layout="wide")

# Page Title
st.markdown("<h1 style='text-align: center; color: #1e88e5;'>📋 Random Question Paper Generator</h1>", unsafe_allow_html=True)

# Input Fields
st.markdown("### 📝 Enter Paper Details")
paper_title = st.text_input("Enter the Heading of the Document")
subject_name = st.text_input("Enter Subject Name")
subject_code = st.text_input("Enter Subject Code")
total_marks = st.number_input("Enter Total Marks", min_value=1, step=1)
time_duration = st.text_input("Enter Time Duration (e.g., 3 Hours)")

# Upload Section with Instructions
st.markdown("### 📂 Upload Your Excel Files")

with st.expander("📌 Important Instructions Before Uploading Excel Files", expanded=True):
    st.markdown("""
    - 🔹 Each Excel file should contain **only one sheet** with the following column headers:
        - `Sr. No`
        - `Questions`
        - `Option 1`
        - `Option 2`
        - `Option 3`
        - `Option 4`
        - `Correct Answer`
    - 🔹 Ensure that all cells under these columns are **filled** and there are **no missing values**.
    - 🔹 File names should be meaningful as the system uses the file name as the **section/topic title**.
    - 🔹 To **exclude questions**, provide their `Sr. No` in the respective exclusion field (comma-separated).
    - 🔹 Avoid using **special characters or merged cells** in your Excel files.
    - 🔹 Each file should contain at least as many questions as you plan to select.
    - 🔹 Maximum file size should not exceed **10MB** per file.
    """)

uploaded_files = st.file_uploader("Upload multiple Excel files", type=['xlsx'], accept_multiple_files=True)

num_questions_per_sheet = {}
marks_per_section = {}
excluded_questions_per_file = {}
used_questions = set()
column_names = ["Sr. No", "Questions", "Option 1", "Option 2", "Option 3", "Option 4", "Correct Answer"]

if uploaded_files:
    st.markdown("### 🔢 Select Number of Questions and Marks Per Section")
    for file in uploaded_files:
        df = pd.read_excel(file)
        if not all(col in df.columns for col in column_names):
            st.error(f"❌ Invalid columns in {file.name}. Expected: {column_names}")
            continue
        max_questions = len(df)
        topic_name = os.path.splitext(file.name)[0]  # Remove file extension
        num_questions_per_sheet[topic_name] = st.number_input(
            f"Questions from {topic_name} (Max: {max_questions})", 1, max_questions, 1, key=topic_name
        )
        marks_per_section[topic_name] = st.number_input(
            f"Marks for {topic_name}", 1, step=1, key=f"marks_{topic_name}"
        )
        excluded_questions_input = st.text_area(
            f"Enter question numbers to exclude for {topic_name} (comma-separated):", key=f"exclude_{topic_name}"
        )
        excluded_questions_per_file[topic_name] = set(excluded_questions_input.split(",")) if excluded_questions_input else set()

# Generate Button
if st.button("🚀 Generate Question Paper"):
    if not uploaded_files:
        st.warning("⚠ Please upload at least one Excel file.")
    else:
        doc = Document()

        # Title Formatting
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run(paper_title.upper())
        title_run.bold = True
        title_run.font.size = Pt(18)

        doc.add_paragraph("\n")  # Space after title

        # Table for Details
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f"Subject: {subject_name}"
        hdr_cells[1].text = f"Subject Code: {subject_code}"
        hdr_cells[2].text = f"Total Marks: {total_marks}"
        hdr_cells[3].text = f"Time: {time_duration}"

        doc.add_paragraph("\n")  # Space after table

        final_data = []
        for file in uploaded_files:
            df = pd.read_excel(file)
            topic_name = os.path.splitext(file.name)[0]
            excluded_set = excluded_questions_per_file[topic_name]
            df = df[~df["Sr. No"].astype(str).isin(excluded_set)]
            if topic_name in num_questions_per_sheet:
                num_questions = num_questions_per_sheet[topic_name]
                df = df[~df["Questions"].isin(used_questions)]
                if len(df) >= num_questions:
                    selected_questions = df.sample(num_questions, random_state=random.randint(1, 100))

                    # Section Title
                    section_title = doc.add_paragraph()
                    section_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    section_run = section_title.add_run(f"Section: {topic_name} (Marks: {marks_per_section[topic_name]})")
                    section_run.bold = True
                    section_run.font.size = Pt(14)

                    final_data.append([topic_name, "", "", "", "", "", ""])  # Topic name row

                    # Reset question counter for this section
                    q_number = 1

                    for _, row in selected_questions.iterrows():
                        # Question
                        question_paragraph = doc.add_paragraph()
                        question_run = question_paragraph.add_run(f"{q_number}. {row['Questions']}")
                        question_run.bold = True
                        question_run.font.size = Pt(12)

                        # Options in one line
                        options_text = f"a) {row['Option 1']}    b) {row['Option 2']}    c) {row['Option 3']}    d) {row['Option 4']}"
                        options_paragraph = doc.add_paragraph(options_text)
                        options_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                        final_data.append(row.tolist())  # Store original row
                        used_questions.add(row['Questions'])
                        q_number += 1
                else:
                    st.warning(f"⚠ Not enough unique questions in {topic_name}. Skipping.")

        # Save Excel File
        if final_data:
            output_df = pd.DataFrame(final_data, columns=column_names)
            excel_buffer = io.BytesIO()
            with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                output_df.to_excel(writer, index=False, sheet_name="Question Paper")
            excel_buffer.seek(0)
            st.session_state["excel_file"] = excel_buffer

        # Save Word File
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        st.session_state["word_file"] = word_buffer
        st.success("✅ Question paper generated successfully!")

# Download Buttons
if "excel_file" in st.session_state and "word_file" in st.session_state:
    st.download_button("📥 Download Question Paper (Excel)", st.session_state["excel_file"], "Generated_Question_Paper.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    st.download_button("📥 Download Question Paper (Word)", st.session_state["word_file"], "Generated_Question_Paper.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
